# Yoinked from https://github.com/twitterdev/Twitter-API-v2-sample-code/blob/main/User-Tweet-Timeline/user_tweets.py

import requests
import os
import json
import docx

# To set your environment variables in your terminal run the following line:
# export 'BEARER_TOKEN'='<your_bearer_token>'
bearer_token = os.environ.get("BEARER_TOKEN")
user_id = 1488106988177440768 # Change this to your user id (https://tweeterid.com/)
testing = False


def create_url():
    return "https://api.twitter.com/2/users/{}/tweets".format(user_id)


def get_params():
    # Tweet fields are adjustable.
    # Options include:
    # attachments, author_id, context_annotations,
    # conversation_id, created_at, entities, geo, id,
    # in_reply_to_user_id, lang, non_public_metrics, organic_metrics,
    # possibly_sensitive, promoted_metrics, public_metrics, referenced_tweets,
    # source, text, and withheld
    return {"tweet.fields": "created_at", "exclude": "retweets,replies", "max_results": "100"}


def bearer_oauth(r):
    """
    Method required by bearer token authentication.
    """

    r.headers["Authorization"] = f"Bearer {bearer_token}"
    r.headers["User-Agent"] = "v2UserTweetsPython"
    return r


def connect_to_endpoint(url, params):
    response = requests.request("GET", url, auth=bearer_oauth, params=params)
    print('API Response: ' + str(response.status_code))
    if response.status_code != 200:
        raise Exception(
            "Request returned an error: {} {}".format(
                response.status_code, response.text
            )
        )
    return response.json()


def main():
    if testing:
        print("Testing mode")
        json_response = read_from_file()
    else:
        print("Live mode")
        url = create_url()
        params = get_params()
        print("Connecting to endpoint...")
        json_response = connect_to_endpoint(url, params)
        print("Writing to json file...")
        write_to_file(json_response)

    # if json response isnt null, write to docx, otherwise print error
    if json_response:
        print("Writing to docx file...")
        write_to_doc(json_response)
    else:
        print("Error: No tweets found/Something went wrong")

# function that writes json response to microsoft word document
# https://python-docx.readthedocs.io/en/latest/user/quickstart.html
def write_to_doc(json_response):
    document = docx.Document()
    document.add_heading('Tweet links')

    for tweet in json_response['data']:
        document.add_paragraph()
        text = 'https://twitter.com/piotrpdev/status/' + str(tweet['id'])
        add_hyperlink(document.add_paragraph(), text, text, None, True)

    document.save('tweets.docx')

# https://stackoverflow.com/a/67447597/10253214
def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

# function that writes json dumps to file
def write_to_file(json_response):
    with open('tweets.json', 'w') as outfile:
        json.dump(json_response, outfile, indent=4)

# function that reads json file and prints the result
def read_from_file():
    with open('tweets.json') as json_file:
        data = json.load(json_file)
        return data


if __name__ == "__main__":
    main()